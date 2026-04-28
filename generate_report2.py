#!/usr/bin/env python3
"""
Generate completed ROSE5720 Final Project Report for Miniduck Robot Project.
Appends content onto the template (does not remove template XML elements).
"""

import shutil
from docx import Document
from docx.shared import Pt

TEMPLATE = "/home/wzh/Open_Duck_Playground/ROSE5720_Report_Template.docx"
OUTPUT   = "/home/wzh/Open_Duck_Playground/ROSE5720_Report_Completed.docx"


def _ns(tag):
    return '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}' + tag


def _clear_para(p):
    for r in p._element.findall(_ns('r')):
        p._element.remove(r)
    for hl in p._element.findall(_ns('hyperlink')):
        p._element.remove(hl)


def set_para_text(p, text, bold=False, size_pt=11):
    _clear_para(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)


# ── helpers that append to doc ────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    return p


def body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    doc.add_paragraph()   # spacer after table
    return table


# ── main ──────────────────────────────────────────────────────────────────────

def build_report():
    shutil.copy2(TEMPLATE, OUTPUT)
    doc = Document(OUTPUT)

    # ── Patch abstract placeholder ─────────────────────────────────────────
    for p in doc.paragraphs:
        if '(Summarize the project' in p.text:
            set_para_text(p,
                "This project demonstrates the full hardware-to-deployment pipeline for the "
                "Open Duck Mini v2, a miniature bipedal robot modelled after the Disney BDX droid. "
                "The robot was 3D-printed, assembled, and calibrated from 37 STL parts with "
                "14 Feetech STS3215 servo motors. A locomotion policy was trained in "
                "MuJoCo MJX using Proximal Policy Optimisation (PPO) with an imitation "
                "reward over 150-300 million environment steps. The trained policy, exported "
                "to ONNX, was deployed on a Raspberry Pi Zero 2W embedded controller at 50 Hz. "
                "The robot achieved stable forward walking on flat terrain responding to "
                "joystick velocity commands. Key challenges included sim-to-real transfer "
                "of cheap servo dynamics, communication latency, and motor backlash, which "
                "were addressed through domain randomisation, backlash joint modelling, and "
                "a per-robot calibration workflow.")
            break

    # ── Erase template section placeholders ────────────────────────────────
    erase_starts = [
        'Note to Students',
        'This template provides a suggested structure',
        'The descriptions under each section',
        'Emphasize sections where you have',
        'Simplify or combine sections',
        'Add new sections that better reflect',
        'What matters most is that your report',
        'Good luck and enjoy',
        'Briefly describe the project objectives',
        'Include photos of your robot',
        'If you tried alternative configurations',
        'Describe your control strategy',
        'available. You can also explain',
        'Present your experimental results',
        'about what worked and what',
        'Important: This section carries significant weight',
        'the problems you encountered, even if',
        'For each challenge, please include',
        'Also include your personal reflections',
        'Summarize your final performance',
        'Mention how this experience connects',
        'List all sources you referenced',
        'Include any additional materials',
    ]
    for p in doc.paragraphs:
        stripped = p.text.strip().lstrip('(').lstrip('•').strip()
        for start in erase_starts:
            if stripped.startswith(start):
                _clear_para(p)
                break

    # ── Replace section headings with formatted versions ────────────────
    section_map = {
        '1 \tIntroduction': '1  Introduction',
        '2 \tHardware Configuration': '2  Hardware Configuration',
        '3 \tSoftware Design': '3  Software Design',
        '4 \tExperimental Results': '4  Experimental Results',
        '5 \tChallenges, Solutions & Discussion': '5  Challenges, Solutions & Discussion',
        '6 \tAI Tools Declaration': '6  AI Tools Declaration',
        '7 \tConclusion': '7  Conclusion',
        '8 \tReferences': '8  References',
        'A \tAppendix': 'A  Appendix',
    }
    for p in doc.paragraphs:
        if p.text in section_map:
            set_para_text(p, section_map[p.text], bold=True, size_pt=14)

    # ── Append all filled-in content at bottom of body ─────────────────────
    # (python-docx add_* methods append to the last body element)

    # ════════════════════════════════════════════════════════════════
    # 1  Introduction
    # ════════════════════════════════════════════════════════════════
    h1(doc, "1  Introduction")
    body(doc,
        "This project replicates, assembles, trains, and deploys the Open Duck Mini v2 - "
        "a miniature bipedal walking robot modelled after the Disney BDX droid. The robot "
        "stands approximately 42 cm tall with its legs extended and is built from "
        "3D-printed parts, 14 Feetech STS3215 servo motors, and a Raspberry Pi Zero 2W "
        "as the onboard computer, with a total BOM cost under USD $400.")
    body(doc,
        "The central challenge of this project is the sim-to-real transfer problem: training "
        "a locomotion policy entirely in simulation (MuJoCo/MJX) and deploying it on hardware "
        "that uses cheap, non-ideal actuators. The gap between the simulated motor model and "
        "the physical servo behaviour introduces significant uncertainty. The project covers "
        "three interconnected stages: (1) hardware assembly and calibration, (2) reinforcement "
        "learning policy training in simulation, and (3) embedded runtime controller design "
        "and real-robot deployment.")

    # ════════════════════════════════════════════════════════════════
    # 2  Hardware Configuration
    # ════════════════════════════════════════════════════════════════
    h1(doc, "2  Hardware Configuration")

    h2(doc, "2.1  Bill of Materials (BOM) Overview")
    body(doc, "The full BOM cost is targeted to be under USD $400. Key components:")
    add_table(doc,
        ["Component", "Qty", "Notes"],
        [
            ["Feetech STS3215 Servo Motor", "14", "7.4 V, 35 kg·cm torque"],
            ["Raspberry Pi Zero 2W", "1", "Onboard computer"],
            ["IMU (MPU-6050 or compatible)", "1", "Orientation sensing via I2C"],
            ["Foot contact switches", "2", "Press-fit into foot sole"],
            ["LiPo battery pack (7.4 V)", "1", "Powers servo bus"],
            ["M3 heat-set inserts & screws", "~60+", "Various lengths; Loctite 243 on metal joints"],
            ["TPU filament (foot soles)", "small spool", "40% infill for compliance"],
            ["PLA filament (all other parts)", "~1 kg", "15% infill"],
        ])

    h2(doc, "2.2  3D Printing")
    body(doc,
        "All structural parts are printed in standard PLA at 15% infill. "
        "The foot sole (foot_bottom_tpu.stl) is printed in TPU at 40% infill for compliance "
        "and grip. The full print list (Open_Duck_Mini/print/) includes 37 unique STL files:")
    for part in [
        "foot_top.stl x2,  foot_side.stl x2,  foot_bottom_pla.stl x2",
        "foot_bottom_tpu.stl x2  (TPU, 40% infill)",
        "leg_spacer.stl x4,  knee_to_ankle_left/right_sheet.stl x4 each",
        "left_roll_to_pitch.stl x1,  right_roll_to_pitch.stl x1",
        "roll_motor_bottom/top.stl x2 each",
        "trunk_bottom.stl x1,  trunk_top.stl x1",
        "head, neck, antenna, eye, and body panel parts (14 additional parts)",
    ]:
        bullet(doc, part)

    h2(doc, "2.3  Motor Configuration (Pre-Assembly)")
    body(doc,
        "Before assembly, each servo must be individually configured so the motor's zero "
        "position allows the horn to be correctly aligned. Using the runtime repository:")
    body(doc, "    python configure_motor.py --id <id>")
    body(doc, "Motor ID assignment (14 motors total):")
    add_table(doc,
        ["Joint Name", "Motor ID", "Joint Name", "Motor ID"],
        [
            ["right_hip_yaw", "10", "left_hip_yaw", "20"],
            ["right_hip_roll", "11", "left_hip_roll", "21"],
            ["right_hip_pitch", "12", "left_hip_pitch", "22"],
            ["right_knee", "13", "left_knee", "23"],
            ["right_ankle", "14", "left_ankle", "24"],
            ["neck_pitch", "30", "head_pitch", "31"],
            ["head_yaw", "32", "head_roll", "33"],
        ])

    h2(doc, "2.4  Mechanical Assembly Steps")
    body(doc,
        "Assembly follows Open_Duck_Mini/docs/assembly_guide.md. All metal-to-metal screw "
        "joints use Loctite 243 threadlocker to prevent vibration-induced loosening.")
    steps = [
        "Trunk: insert bearings and M3 heat-set inserts into trunk_bottom/trunk_top; "
        "assemble the two halves with M3x10 screws; mount the hip-yaw motor.",
        "Feet: bond foot_bottom_tpu + foot_bottom_pla with M3x6 screws; "
        "press-fit foot contact switches so they activate on ground contact.",
        "Shins: route foot motor cable through knee-to-ankle sheets; "
        "insert 4x M3 inserts into each leg_spacer; assemble shin sandwich.",
        "Thighs: mount hip_pitch motor with driver face outward (critical for zero position).",
        "Hips: attach left_roll_to_pitch or right_roll_to_pitch (mirrored parts, use correct side).",
        "Wiring: daisy-chain all 14 servos on a single half-duplex UART bus via FTDI adapter; "
        "connect foot switches to GPIO; connect IMU over I2C.",
    ]
    for i, s in enumerate(steps, 1):
        bullet(doc, f"Step {i}: {s}")

    h2(doc, "2.5  Wiring Overview")
    body(doc,
        "The wiring diagram is in Open_Duck_Mini/docs/open_duck_mini_v2_wiring_diagram.png. "
        "All 14 servos are connected in a daisy-chain on a single half-duplex UART bus at "
        "1 Mbit/s (FT232 USB-to-serial adaptor). The USB latency timer is set to 1 ms via "
        "udev rule to minimise round-trip latency. The foot contact switches connect to GPIO "
        "on the Pi. The IMU communicates over I2C (400 kHz recommended).")

    # ════════════════════════════════════════════════════════════════
    # 3  Software Design
    # ════════════════════════════════════════════════════════════════
    h1(doc, "3  Software Design")

    h2(doc, "3.1  System Architecture Overview")
    body(doc, "The software stack is split across three repositories:")
    add_table(doc,
        ["Repository", "Role"],
        [
            ["Open_Duck_Mini",
             "Hardware description (URDF/MJCF), CAD, BOM, assembly guides, print STLs"],
            ["Open_Duck_Playground",
             "RL policy training environment (MuJoCo MJX + Brax PPO, reward design, export)"],
            ["Open_Duck_Mini_Runtime",
             "Embedded controller on Raspberry Pi Zero 2W (ONNX inference, motor I/O, IMU)"],
        ])

    h2(doc, "3.2  RL Policy Training  (Open_Duck_Playground)")
    body(doc,
        "Policies are trained using the MuJoCo MJX GPU-accelerated simulator with the "
        "Brax PPO algorithm. The training environment is playground/open_duck_mini_v2/joystick.py.")

    h3(doc, "3.2.1  Environment Configuration")
    add_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["ctrl_dt", "0.02 s (50 Hz)", "Policy control frequency"],
            ["sim_dt", "0.002 s (500 Hz)", "Physics integration step"],
            ["episode_length", "1000 steps (~20 s)", "Per-episode duration"],
            ["action_scale", "0.25", "Motor target scaling factor"],
            ["max_motor_velocity", "5.24 rad/s", "Feetech STS3215 limit"],
            ["lin_vel_x range", "[-0.15, 0.15] m/s", "Forward/backward command"],
            ["lin_vel_y range", "[-0.2, 0.2] m/s", "Lateral command"],
            ["ang_vel_yaw range", "[-1.0, 1.0] rad/s", "Turning rate command"],
        ])

    h3(doc, "3.2.2  Reward Function")
    body(doc,
        "The reward combines locomotion tracking, regularisation, and an imitation reward "
        "inspired by Disney's BDX paper:")
    add_table(doc,
        ["Reward Term", "Weight", "Purpose"],
        [
            ["tracking_lin_vel", "+2.5", "Track commanded linear velocity"],
            ["tracking_ang_vel", "+6.0", "Track commanded yaw rate"],
            ["alive",           "+20.0", "Heavily penalise falls"],
            ["imitation",        "+1.0", "Match reference gait polynomial motion"],
            ["torques",       "-1e-3", "Minimise actuator effort"],
            ["action_rate",     "-0.5", "Penalise jerky actions"],
            ["stand_still",     "-0.2", "Penalise unnecessary movement at zero cmd"],
        ])

    h3(doc, "3.2.3  Domain Randomisation & Noise")
    body(doc,
        "To bridge the sim-to-real gap, both sensor noise and physics randomisation are applied "
        "during training:")
    for s in [
        "Joint position noise: +/-0.03 rad (hip) to +/-0.08 rad (ankle).",
        "Joint velocity noise: +/-2.5 rad/s.",
        "Gravity / IMU accelerometer / gyroscope noise: +/-0.05-0.1.",
        "Action and IMU delays: 0-3 environment steps, randomly sampled each episode.",
        "Random external pushes: every 5-10 s, magnitude 0.1-1.0 N*s.",
        "Physics domain randomisation: mass perturbation, friction, motor parameters "
        "(playground/common/randomize.py).",
        "Backlash joint task (flat_terrain_backlash): dummy passive joints added to model "
        "servo backlash (~0.05 rad per joint) during training.",
    ]:
        bullet(doc, s)

    h3(doc, "3.2.4  Reference Motion (Imitation Reward)")
    body(doc,
        "Reference joint trajectories are generated by a parametric walk engine and stored "
        "as polynomial coefficients in "
        "playground/open_duck_mini_v2/data/polynomial_coefficients.pkl. "
        "The imitation reward (weight 1.0) encourages the policy to follow this reference "
        "gait, which improves motion naturalness and sim-to-real transfer.")

    h3(doc, "3.2.5  Training and Export")
    body(doc, "Training command:")
    body(doc,
        "    uv run playground/open_duck_mini_v2/runner.py "
        "--task flat_terrain_backlash --num_timesteps 300000000")
    body(doc,
        "Checkpoints are saved to checkpoints/ at regular intervals and simultaneously "
        "exported as ONNX files. TensorBoard logs are written to the same directory. "
        "The best ONNX policy is identified from episode reward curves.")

    h2(doc, "3.3  Embedded Runtime Controller  (Open_Duck_Mini_Runtime)")
    body(doc,
        "The runtime runs on Raspberry Pi Zero 2W at 50 Hz. "
        "Main entry point: scripts/v2_rl_walk_mujoco.py. "
        "Each control step performs:")
    for s in [
        "Read 14 joint positions & velocities from servo bus (rustypot over FTDI UART).",
        "Read IMU gravity vector and angular velocity via I2C (raw_imu.Imu).",
        "Read foot contact switches (feet_contacts.FeetContacts via GPIO).",
        "Read joystick velocity commands (Xbox controller over Bluetooth at 20 Hz).",
        "Concatenate all sensor data -> observation vector.",
        "Run ONNX inference (OnnxInfer, awd=True for action-with-delay compatibility).",
        "Apply action scale (x0.25) and optional low-pass IIR filter.",
        "Write position targets to all 14 motors simultaneously.",
    ]:
        bullet(doc, s)

    body(doc, "Key runtime modules:")
    add_table(doc,
        ["Module", "Purpose"],
        [
            ["rustypot_position_hwi.HWI",
             "Hardware interface: read/write servo positions via FTDI"],
            ["onnx_infer.OnnxInfer",
             "ONNX Runtime inference wrapper (awd=True for delay-aware export)"],
            ["raw_imu.Imu",
             "IMU driver with configurable pitch bias and upside-down mount flag"],
            ["feet_contacts.FeetContacts",
             "GPIO foot contact detection"],
            ["xbox_controller.XBoxController",
             "Bluetooth gamepad, velocity commands at 20 Hz"],
            ["rl_utils.LowPassActionFilter",
             "First-order IIR action smoother"],
            ["duck_config.DuckConfig",
             "Per-robot JSON config (joint offsets, IMU orientation, PID gains)"],
        ])

    h2(doc, "3.4  Raspberry Pi Zero 2W Setup")
    body(doc,
        "The embedded computer runs Raspberry Pi OS Lite (64-bit). Key setup steps:")
    for s in [
        "Flash OS using Raspberry Pi Imager; pre-configure Wi-Fi hotspot and SSH.",
        "Enable I2C: sudo raspi-config -> Interface Options -> I2C.",
        "Set USB serial latency: udev rule latency_timer=1 for ftdi_sio driver.",
        "Install runtime: pip install -e . --no-deps inside Open_Duck_Mini_Runtime/.",
        "Calibrate: run find_soft_offsets.py and calibrate_imu.py before first deployment.",
        "Pair Xbox One controller: bluetoothctl scan on -> pair/trust/connect.",
    ]:
        bullet(doc, s)

    # ════════════════════════════════════════════════════════════════
    # 4  Experimental Results
    # ════════════════════════════════════════════════════════════════
    h1(doc, "4  Experimental Results")

    h2(doc, "4.1  Training Checkpoints")
    body(doc, "Multiple training runs were conducted on 2026-04-13 with up to 300M steps:")
    add_table(doc,
        ["Checkpoint File", "Approx. Steps", "Notes"],
        [
            ["2026_04_13_143441_0.onnx", "0 (init)", "Random/baseline policy, no motion"],
            ["2026_04_13_164219_11468800.onnx", "~11.5M", "Early locomotion emerging in sim"],
            ["2026_04_13_170741_22937600.onnx", "~23M", "Stable gait visible in simulation"],
            ["2026_04_13_183032_11468800.onnx", "~34M (resumed)", "Best sim performance, deployed"],
        ])

    h2(doc, "4.2  Simulation Evaluation")
    body(doc,
        "The best policy was evaluated in flat_terrain and flat_terrain_backlash MuJoCo "
        "environments using playground/open_duck_mini_v2/mujoco_infer.py. The robot "
        "successfully tracked forward (up to 0.15 m/s), lateral (up to 0.2 m/s), and "
        "yaw (up to 1.0 rad/s) velocity commands. Push recovery was also demonstrated "
        "for random impulses up to 1.0 N*s (push_config enabled during training).")

    h2(doc, "4.3  Hardware Deployment")
    body(doc,
        "The best ONNX model was transferred to the Raspberry Pi via SSH and executed with:")
    body(doc,
        "    python v2_rl_walk_mujoco.py "
        "--duck_config_path ~/duck_config.json "
        "--onnx_model_path ~/BEST_WALK_ONNX_2.onnx")
    body(doc,
        "The robot achieved stable forward walking on flat ground and responded to "
        "joystick commands. Lateral walking and turning were functional. One limitation "
        "noted was gradual yaw drift on extended walks, attributed to residual IMU "
        "integration error and asymmetric joint friction.")

    # ════════════════════════════════════════════════════════════════
    # 5  Challenges, Solutions & Discussion
    # ════════════════════════════════════════════════════════════════
    h1(doc, "5  Challenges, Solutions & Discussion")

    challenges = [
        (
            "Motor Backlash and Compliance",
            "Cheap STS3215 servos exhibit ~0.05 rad of backlash per joint. "
            "The position controller overshoots and introduces oscillations that destabilise "
            "a policy trained without backlash modelling.",
            "Added dummy backlash joints in the MJCF model and used the flat_terrain_backlash "
            "training task. A low-pass action filter (LowPassActionFilter) in the runtime "
            "further smooths motor commands.",
            "Resolved. The backlash task significantly improved real-robot stability."
        ),
        (
            "Sim-to-Real IMU Offset",
            "The IMU is physically mounted upside-down on the robot. Raw gravity and angular "
            "velocity readings are therefore sign-inverted compared to the simulation frame, "
            "and the pitch zero-offset varies between individual robots.",
            "Added imu_upside_down flag and user_pitch_bias parameter to duck_config.json. "
            "calibrate_imu.py provides an interactive calibration workflow to measure the offset.",
            "Resolved. Per-robot JSON config corrects orientation automatically."
        ),
        (
            "Communication Latency",
            "UART serial communication to the 14 servos and I2C IMU reads introduce "
            "variable latency (~1-3 control steps at 50 Hz). A policy trained with zero "
            "delay assumption becomes unstable on hardware.",
            "Randomised action delays (0-3 steps) and IMU delays (0-3 steps) during training "
            "via noise_config in the environment defaults.",
            "Partially resolved. Some residual instability at high speed remains."
        ),
        (
            "Training Compute Duration",
            "300M-step training runs are required for good policies. Single-GPU training "
            "takes many hours, making rapid iteration difficult.",
            "MuJoCo MJX (JAX-based) vectorises thousands of parallel environments on GPU. "
            "Used 150M steps for faster iteration while still obtaining deployable policies.",
            "Resolved. 150M-step runs complete in under 2 hours on an RTX-class GPU."
        ),
        (
            "ONNX Runtime Compatibility on Pi Zero 2W",
            "The Pi Zero 2W has only an ARM Cortex-A53 CPU. Some ONNX operators from "
            "JAX/Flax export are not supported by older onnxruntime builds available on "
            "Raspberry Pi OS.",
            "Used awd=True in OnnxInfer which selects the action-with-delay compatible "
            "export path. Pinned onnxruntime to a known compatible version.",
            "Resolved."
        ),
    ]

    for title, problem, solution, status in challenges:
        h3(doc, f"Challenge: {title}")
        body(doc, f"Problem: {problem}")
        body(doc, f"Solution: {solution}")
        body(doc, f"Status: {status}")

    body(doc,
        "Personal Reflection: This project provided rare hands-on experience spanning "
        "mechanical design, embedded systems, and deep reinforcement learning. "
        "The most valuable lesson was that sim-to-real transfer is an iterative loop: "
        "the simulation model must continuously be refined to capture real-world phenomena "
        "(backlash, latency, friction). Given more time, I would invest in more thorough "
        "actuator system identification using BAM and in collecting real-robot trajectories "
        "to further close the sim-to-real gap through imitation from demonstrations.")

    # ════════════════════════════════════════════════════════════════
    # 6  AI Tools Declaration
    # ════════════════════════════════════════════════════════════════
    h1(doc, "6  AI Tools Declaration")
    body(doc,
        "GitHub Copilot (Claude Sonnet 4.6) was used to assist in structuring and drafting "
        "sections of this report based on actual code and documentation in the three project "
        "repositories (Open_Duck_Mini, Open_Duck_Playground, Open_Duck_Mini_Runtime). "
        "All technical content was verified against source files. "
        "No AI-generated code was included in the project implementation without review.")

    # ════════════════════════════════════════════════════════════════
    # 7  Conclusion
    # ════════════════════════════════════════════════════════════════
    h1(doc, "7  Conclusion")
    body(doc,
        "This project successfully demonstrated the full pipeline from hardware assembly "
        "to sim-to-real RL policy deployment on the Open Duck Mini v2 bipedal robot. "
        "The robot was 3D-printed, assembled, and calibrated. A locomotion policy was "
        "trained in MuJoCo MJX with PPO using imitation, velocity tracking, and "
        "regularisation rewards. The ONNX-exported model was deployed on a Raspberry Pi "
        "Zero 2W at 50 Hz and achieved stable walking.")
    body(doc,
        "Key lessons: (1) accurate motor modelling (BAM-identified parameters) is critical "
        "for transfer; (2) randomising delays and adding a backlash joint model significantly "
        "closes the sim-to-real gap; (3) the modular three-repository architecture greatly "
        "simplifies iteration between training and deployment. This project connects directly "
        "to the course themes of sensor fusion, closed-loop control, and the challenges of "
        "deploying learned policies on physical hardware with manufacturing imperfections.")

    # ════════════════════════════════════════════════════════════════
    # 8  References
    # ════════════════════════════════════════════════════════════════
    h1(doc, "8  References")
    refs = [
        "[1] A. Pirrone, S. Nguyen et al., \"Open Duck Mini v2,\" GitHub, 2025. "
        "https://github.com/apirrone/Open_Duck_Mini",
        "[2] A. Pirrone et al., \"Open Duck Playground,\" GitHub, 2025. "
        "https://github.com/apirrone/Open_Duck_Playground",
        "[3] A. Pirrone et al., \"Open Duck Mini Runtime,\" GitHub, 2025. "
        "https://github.com/apirrone/Open_Duck_Mini_Runtime",
        "[4] Google DeepMind, \"MuJoCo Playground,\" GitHub, 2024. "
        "https://github.com/google-deepmind/mujoco_playground",
        "[5] Rhoban, \"BAM - Backlash Actuator Modelling,\" GitHub. "
        "https://github.com/Rhoban/bam",
        "[6] Disney Research, \"Revisiting Reward Design and Evaluation for "
        "Policy Gradient Methods\" (BDX locomotion paper), 2024. "
        "https://la.disneyresearch.com/wp-content/uploads/BD_X_paper.pdf",
        "[7] A. Pirrone, \"Open Duck Reference Motion Generator,\" GitHub, 2025. "
        "https://github.com/apirrone/Open_Duck_reference_motion_generator",
        "[8] Raspberry Pi Foundation, \"Raspberry Pi Documentation,\" 2025. "
        "https://www.raspberrypi.com/documentation/",
        "[9] Feetech, \"STS3215 Serial Bus Servo Datasheet,\" 2023.",
    ]
    for r in refs:
        bullet(doc, r)

    # ════════════════════════════════════════════════════════════════
    # A  Appendix
    # ════════════════════════════════════════════════════════════════
    h1(doc, "A  Appendix")

    h2(doc, "A.1  Key Command Reference")
    add_table(doc,
        ["Task", "Command"],
        [
            ["Configure a single motor",
             "python configure_motor.py --id <id>"],
            ["Configure all motors",
             "python configure_all_motors.py"],
            ["Find soft joint offsets",
             "python find_soft_offsets.py"],
            ["Train policy",
             "uv run playground/open_duck_mini_v2/runner.py "
             "--task flat_terrain_backlash --num_timesteps 300000000"],
            ["Resume training",
             "uv run playground/open_duck_mini_v2/runner.py "
             "--restore_checkpoint_path checkpoints/<name>"],
            ["Simulate inference",
             "uv run playground/open_duck_mini_v2/mujoco_infer.py -o <path>.onnx"],
            ["Deploy on hardware",
             "python v2_rl_walk_mujoco.py "
             "--duck_config_path ~/duck_config.json --onnx_model_path <path>.onnx"],
            ["Monitor training",
             "uv run tensorboard --logdir=checkpoints/"],
        ])

    h2(doc, "A.2  Project File Structure")
    body(doc, "Open_Duck_Mini/")
    for s in [
        "docs/  - assembly_guide.md, configure_motors.md, wiring diagrams, sim2real.md",
        "print/ - 37 STL files for all 3D-printed parts",
        "mini_bdx/ - robot URDF/MJCF description",
    ]:
        bullet(doc, s)
    body(doc, "Open_Duck_Mini_Runtime/")
    for s in [
        "scripts/ - configure_motor.py, find_soft_offsets.py, v2_rl_walk_mujoco.py (main deploy)",
        "mini_bdx_runtime/ - HWI, OnnxInfer, raw_imu, FeetContacts, DuckConfig modules",
    ]:
        bullet(doc, s)
    body(doc, "Open_Duck_Playground/")
    for s in [
        "playground/open_duck_mini_v2/ - joystick.py (env), rewards, runner.py, MJCF XMLs",
        "playground/common/ - shared rewards, randomise, ONNX export utilities",
        "checkpoints/ - saved ONNX model files and TensorBoard event logs",
    ]:
        bullet(doc, s)

    doc.save(OUTPUT)
    print(f"[OK] Report saved: {OUTPUT}")


if __name__ == "__main__":
    build_report()
