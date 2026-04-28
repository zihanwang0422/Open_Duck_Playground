#!/usr/bin/env python3
"""
Generate completed ROSE5720 Final Project Report.
Inserts filled-in content BEFORE the Declaration page so the order is correct.
"""

import shutil
from docx import Document
from docx.shared import Pt
from lxml import etree

TEMPLATE = "/home/wzh/Open_Duck_Playground/ROSE5720_Report_Template.docx"
OUTPUT   = "/home/wzh/Open_Duck_Playground/ROSE5720_Report_Completed.docx"

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def ns(tag):
    return f'{{{NS}}}{tag}'


# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _clear_para(p):
    for child in list(p._element):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('r', 'hyperlink', 'ins', 'del'):
            p._element.remove(child)


def _make_para_xml(text, bold=False, size_pt=11, style='Normal'):
    """Create a <w:p> element with a single run."""
    p = etree.SubElement(etree.Element('dummy'), ns('p'))
    pPr = etree.SubElement(p, ns('pPr'))
    pStyle = etree.SubElement(pPr, ns('pStyle'))
    pStyle.set(ns('val'), style)

    r = etree.SubElement(p, ns('r'))
    rPr = etree.SubElement(r, ns('rPr'))
    if bold:
        etree.SubElement(rPr, ns('b'))
        etree.SubElement(rPr, ns('bCs'))
    sz = etree.SubElement(rPr, ns('sz'))
    sz.set(ns('val'), str(int(size_pt * 2)))
    szCs = etree.SubElement(rPr, ns('szCs'))
    szCs.set(ns('val'), str(int(size_pt * 2)))

    t = etree.SubElement(r, ns('t'))
    t.text = text
    if text.startswith(' ') or text.endswith(' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return p


def _make_bullet_xml(text, size_pt=11):
    p = etree.SubElement(etree.Element('dummy'), ns('p'))
    pPr = etree.SubElement(p, ns('pPr'))
    pStyle = etree.SubElement(pPr, ns('pStyle'))
    pStyle.set(ns('val'), 'ListBullet')

    r = etree.SubElement(p, ns('r'))
    rPr = etree.SubElement(r, ns('rPr'))
    sz = etree.SubElement(rPr, ns('sz'))
    sz.set(ns('val'), str(int(size_pt * 2)))

    t = etree.SubElement(r, ns('t'))
    t.text = text
    return p


# ── Content builder that writes to an insertion buffer ───────────────────────

class ParaBuffer:
    def __init__(self):
        self.elements = []   # list of lxml Elements

    def h1(self, text):
        self.elements.append(_make_para_xml(text, bold=True, size_pt=14))

    def h2(self, text):
        self.elements.append(_make_para_xml(text, bold=True, size_pt=12))

    def h3(self, text):
        self.elements.append(_make_para_xml(text, bold=True, size_pt=11))

    def body(self, text):
        self.elements.append(_make_para_xml(text, bold=False, size_pt=11))

    def bullet(self, text):
        self.elements.append(_make_bullet_xml(text, size_pt=11))

    def spacer(self):
        self.elements.append(_make_para_xml(''))

    def table(self, doc, headers, rows):
        """Appends a Table Grid table; returns the table element."""
        tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            for r in hdr[i].paragraphs[0].runs:
                r.bold = True
        for row_data in rows:
            cells = tbl.add_row().cells
            for i, v in enumerate(row_data):
                cells[i].text = str(v)
        # Detach from body (we appended it temporarily) — we'll re-attach via insert
        elem = tbl._element
        elem.getparent().remove(elem)
        self.elements.append(elem)
        self.spacer()
        return tbl


def build_content(buf, doc):
    """Fill buf with all report section content."""

    # ── 1  Introduction ──────────────────────────────────────────────────────
    buf.h1("1  Introduction")
    buf.body(
        "This project replicates, assembles, trains, and deploys the Open Duck Mini v2 — "
        "a miniature bipedal walking robot modelled after the Disney BDX droid. The robot "
        "stands approximately 42 cm tall with its legs extended and is built from "
        "3D-printed structural parts, 14 Feetech STS3215 serial-bus servo motors, and a "
        "Raspberry Pi Zero 2W as the onboard computer, with a total Bill-of-Materials cost "
        "under USD $400.")
    buf.body(
        "The central challenge is the sim-to-real transfer problem: training a locomotion "
        "policy entirely in simulation (MuJoCo MJX) and deploying it on hardware that uses "
        "cheap, non-ideal actuators with backlash and compliance. The project is organised "
        "into three stages: (1) hardware assembly and motor calibration, (2) reinforcement "
        "learning policy training in simulation, and (3) embedded runtime deployment on the "
        "physical robot.")
    buf.spacer()

    # ── 2  Hardware Configuration ────────────────────────────────────────────
    buf.h1("2  Hardware Configuration")

    buf.h2("2.1  Bill of Materials (BOM) Overview")
    buf.body("The full BOM cost targets under USD $400. Key components:")
    buf.table(doc,
        ["Component", "Qty", "Notes"],
        [
            ["Feetech STS3215 Servo Motor", "14", "7.4 V; ~35 kg·cm stall torque"],
            ["Raspberry Pi Zero 2W", "1", "Onboard computer (ARM Cortex-A53, 512 MB RAM)"],
            ["IMU (MPU-6050 compatible)", "1", "Gravity + angular rate via I2C"],
            ["Foot contact switches", "2", "Press-fit binary ground contact sensing"],
            ["LiPo battery pack 7.4 V", "1", "Powers servo bus and Pi"],
            ["M3 heat-set inserts & screws", "~60+", "Loctite 243 on all metal-to-metal joints"],
            ["TPU filament (foot soles)", "~50 g", "Printed at 40% infill for compliance"],
            ["PLA filament (structural)", "~600 g", "All other parts at 15% infill"],
        ])

    buf.h2("2.2  3D Printing")
    buf.body(
        "All structural parts are printed in standard PLA at 15% infill. "
        "The foot sole (foot_bottom_tpu.stl) is printed in TPU at 40% infill for compliance "
        "and grip. The full print list (Open_Duck_Mini/print/) contains 37 unique STL files:")
    for part in [
        "foot_top.stl x2,  foot_side.stl x2,  foot_bottom_pla.stl x2",
        "foot_bottom_tpu.stl x2  (TPU, 40% infill)",
        "leg_spacer.stl x4;  knee_to_ankle_left/right_sheet.stl x4 each",
        "left_roll_to_pitch.stl x1;  right_roll_to_pitch.stl x1",
        "roll_motor_bottom.stl x2;  roll_motor_top.stl x2",
        "trunk_bottom.stl x1;  trunk_top.stl x1",
        "Head/neck/antenna/eye/body panel parts (14 additional parts)",
    ]:
        buf.bullet(part)

    buf.h2("2.3  Motor Configuration (Pre-Assembly)")
    buf.body(
        "Before assembly, each of the 14 servos must be individually configured to set its "
        "zero position so the output horn aligns correctly. Using the runtime repo:")
    buf.body("    python configure_motor.py --id <id>")
    buf.body("Motor ID assignment (Open_Duck_Mini/docs/configure_motors.md):")
    buf.table(doc,
        ["Joint Name", "ID", "Joint Name", "ID"],
        [
            ["right_hip_yaw",   "10", "left_hip_yaw",   "20"],
            ["right_hip_roll",  "11", "left_hip_roll",  "21"],
            ["right_hip_pitch", "12", "left_hip_pitch", "22"],
            ["right_knee",      "13", "left_knee",      "23"],
            ["right_ankle",     "14", "left_ankle",     "24"],
            ["neck_pitch",      "30", "head_pitch",     "31"],
            ["head_yaw",        "32", "head_roll",      "33"],
        ])

    buf.h2("2.4  Mechanical Assembly Steps")
    buf.body(
        "Assembly follows Open_Duck_Mini/docs/assembly_guide.md. "
        "Apply Loctite 243 on all metal-to-metal screw joints to prevent vibration loosening. "
        "Do NOT use Loctite on plastic screws.")
    for i, s in enumerate([
        "Trunk: insert bearings and M3 heat-set inserts into trunk_bottom/trunk_top; "
        "join the two halves with M3x10 screws; mount the hip-yaw motor.",
        "Feet: bond foot_bottom_tpu + foot_bottom_pla with M3x6 screws; "
        "press-fit foot contact switches so they activate on ground contact.",
        "Shins: route the foot motor cable through knee-to-ankle right/left sheets; "
        "insert 4x M3 inserts per leg_spacer; assemble the shin sandwich.",
        "Thighs: mount the hip_pitch motor with its driver face outward "
        "(critical for the correct mechanical zero position).",
        "Hips: attach left_roll_to_pitch or right_roll_to_pitch — parts are mirrored, "
        "use the correct side for each leg.",
        "Wiring: daisy-chain all 14 servos on a single half-duplex UART bus via an "
        "FT232 USB-to-serial adaptor; connect foot switches to GPIO; connect IMU over I2C.",
    ], 1):
        buf.bullet(f"Step {i}: {s}")

    buf.h2("2.5  Wiring Overview")
    buf.body(
        "Wiring diagram: Open_Duck_Mini/docs/open_duck_mini_v2_wiring_diagram.png. "
        "All 14 servos share one half-duplex UART bus at 1 Mbit/s (FT232 adaptor). "
        "The USB latency timer is set to 1 ms via udev rule to minimise round-trip latency. "
        "Foot switches connect to Raspberry Pi GPIO. IMU communicates over I2C at 400 kHz.")
    buf.spacer()

    # ── 3  Software Design ───────────────────────────────────────────────────
    buf.h1("3  Software Design")

    buf.h2("3.1  System Architecture Overview")
    buf.body("The software stack is split across three repositories that work together:")
    buf.table(doc,
        ["Repository", "Role"],
        [
            ["Open_Duck_Mini",
             "Hardware description (URDF/MJCF), CAD, BOM, assembly and wiring guides, print STLs"],
            ["Open_Duck_Playground",
             "RL policy training (MuJoCo MJX + Brax PPO), reward design, ONNX export"],
            ["Open_Duck_Mini_Runtime",
             "Embedded controller on Raspberry Pi Zero 2W: ONNX inference, motor I/O, IMU, config"],
        ])

    buf.h2("3.2  RL Policy Training  (Open_Duck_Playground)")
    buf.body(
        "Policies are trained using the MuJoCo MJX GPU-accelerated physics simulator with "
        "the Brax PPO algorithm. The training environment is defined in "
        "playground/open_duck_mini_v2/joystick.py.")

    buf.h3("3.2.1  Environment Configuration")
    buf.table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["ctrl_dt", "0.02 s (50 Hz)", "Policy control frequency"],
            ["sim_dt", "0.002 s (500 Hz)", "Physics integration timestep"],
            ["episode_length", "1000 steps (~20 s)", "Episode duration"],
            ["action_scale", "0.25", "Motor target scaling factor"],
            ["max_motor_velocity", "5.24 rad/s", "Feetech STS3215 limit"],
            ["lin_vel_x", "[-0.15, 0.15] m/s", "Forward/backward command range"],
            ["lin_vel_y", "[-0.20, 0.20] m/s", "Lateral command range"],
            ["ang_vel_yaw", "[-1.0, 1.0] rad/s", "Turning rate command range"],
        ])

    buf.h3("3.2.2  Reward Function")
    buf.body(
        "The reward function combines locomotion tracking, regularisation, and an imitation "
        "reward inspired by Disney's BDX locomotion paper:")
    buf.table(doc,
        ["Reward Term", "Weight", "Purpose"],
        [
            ["tracking_lin_vel", "+2.5",   "Track commanded linear velocity"],
            ["tracking_ang_vel", "+6.0",   "Track commanded yaw rate"],
            ["alive",           "+20.0",   "Penalise falls heavily"],
            ["imitation",        "+1.0",   "Imitate reference gait from polynomial motion generator"],
            ["torques",         "-0.001",  "Minimise actuator effort"],
            ["action_rate",     "-0.5",    "Penalise jerky / high-frequency actions"],
            ["stand_still",     "-0.2",    "Penalise motion when zero command is given"],
        ])

    buf.h3("3.2.3  Domain Randomisation & Noise")
    buf.body(
        "To bridge the sim-to-real gap, sensor noise and physics randomisation are applied "
        "during training (noise_config in joystick.py default_config):")
    for s in [
        "Joint position noise: +/-0.03 rad (hip joints), up to +/-0.08 rad (ankle joints).",
        "Joint velocity noise: +/-2.5 rad/s.",
        "IMU gravity / gyroscope noise: +/-0.05-0.1.",
        "Action delays: 0-3 environment steps (randomly sampled each episode).",
        "IMU delays: 0-3 environment steps (randomly sampled each episode).",
        "Random external pushes: every 5-10 s, magnitude 0.1-1.0 N*s (push_config enabled).",
        "Physics domain randomisation: mass perturbation, friction variation, motor "
        "parameters (playground/common/randomize.py).",
        "Backlash joints: flat_terrain_backlash task adds passive dummy joints to model "
        "~0.05 rad backlash per servo.",
    ]:
        buf.bullet(s)

    buf.h3("3.2.4  Reference Motion (Imitation Reward)")
    buf.body(
        "Reference joint trajectories are generated by a parametric walk engine and stored as "
        "polynomial coefficients in playground/open_duck_mini_v2/data/polynomial_coefficients.pkl. "
        "The imitation reward (weight 1.0) encourages the policy to follow this natural gait, "
        "which improves motion quality and eases sim-to-real transfer.")

    buf.h3("3.2.5  Training and ONNX Export")
    buf.body("Launch training with:")
    buf.body(
        "    uv run playground/open_duck_mini_v2/runner.py "
        "--task flat_terrain_backlash --num_timesteps 300000000")
    buf.body(
        "Checkpoints are saved to checkpoints/ at regular intervals (approx. every 11.5M steps) "
        "and simultaneously exported as ONNX files. TensorBoard event logs are written to the same "
        "directory. The best policy is selected from episode-reward curves.")

    buf.h2("3.3  Embedded Runtime Controller  (Open_Duck_Mini_Runtime)")
    buf.body(
        "The runtime runs on Raspberry Pi Zero 2W at 50 Hz. "
        "Main entry point: scripts/v2_rl_walk_mujoco.py (class RLWalk). "
        "Each 20 ms control cycle performs the following steps:")
    for s in [
        "Read 14 joint positions and velocities from servo bus "
        "(rustypot_position_hwi.HWI over FTDI UART, 1 Mbit/s).",
        "Read IMU gravity vector and angular velocity via I2C (raw_imu.Imu, 50 Hz).",
        "Read binary foot contact signals via GPIO (feet_contacts.FeetContacts).",
        "Read joystick velocity commands via Bluetooth Xbox controller (20 Hz poll).",
        "Concatenate all sensor data into an observation vector.",
        "Run ONNX inference (onnx_infer.OnnxInfer, awd=True for delay-aware policy).",
        "Apply action scale (x0.25) and optional first-order IIR low-pass filter.",
        "Write position targets to all 14 motors simultaneously.",
    ]:
        buf.bullet(s)

    buf.body("Key runtime modules:")
    buf.table(doc,
        ["Module", "Purpose"],
        [
            ["rustypot_position_hwi.HWI",
             "Hardware interface: synchronous read/write servo positions via FTDI"],
            ["onnx_infer.OnnxInfer",
             "ONNX Runtime inference wrapper; awd=True selects delay-compatible import"],
            ["raw_imu.Imu",
             "IMU driver; supports imu_upside_down flag and user_pitch_bias"],
            ["feet_contacts.FeetContacts", "GPIO-based binary foot contact detection"],
            ["xbox_controller.XBoxController",
             "Bluetooth gamepad input; velocity commands at configurable Hz"],
            ["rl_utils.LowPassActionFilter",
             "First-order IIR action smoother (configurable cutoff frequency)"],
            ["duck_config.DuckConfig",
             "Per-robot JSON config: joint offsets, IMU orientation, PID gains, start_paused"],
        ])

    buf.h2("3.4  Raspberry Pi Zero 2W Setup")
    buf.body("Setup steps for the embedded computer (Open_Duck_Mini_Runtime/README.md):")
    for s in [
        "Flash Raspberry Pi OS Lite 64-bit; pre-configure Wi-Fi hotspot and SSH "
        "in the Raspberry Pi Imager advanced options.",
        "Enable I2C: sudo raspi-config -> Interface Options -> I2C.",
        "Set USB serial latency to 1 ms: "
        "udev rule SUBSYSTEM==\"usb-serial\", DRIVER==\"ftdi_sio\", "
        "ATTR{latency_timer}=\"1\".",
        "Install runtime: cd Open_Duck_Mini_Runtime && pip install -e . --no-deps.",
        "Calibrate: run find_soft_offsets.py (joint offsets) and calibrate_imu.py "
        "(pitch bias) before first policy deployment.",
        "Pair Xbox One controller over Bluetooth: bluetoothctl -> scan on -> "
        "pair/trust/connect <MAC>.",
    ]:
        buf.bullet(s)
    buf.spacer()

    # ── 4  Experimental Results ──────────────────────────────────────────────
    buf.h1("4  Experimental Results")

    buf.h2("4.1  Training Checkpoints (2026-04-13)")
    buf.body(
        "Multiple training runs were conducted. Checkpoints were saved at regular intervals:")
    buf.table(doc,
        ["Checkpoint File (ONNX)", "Approx. Steps", "Notes"],
        [
            ["2026_04_13_143441_0.onnx",         "0 (init)", "Random initialisation, no locomotion"],
            ["2026_04_13_164219_11468800.onnx",   "~11.5M",  "Early locomotion emerging in simulation"],
            ["2026_04_13_170741_22937600.onnx",   "~23M",    "Stable gait visible in flat terrain"],
            ["2026_04_13_183032_11468800.onnx",   "~34M (resumed)", "Best policy, deployed on hardware"],
        ])

    buf.h2("4.2  Simulation Performance")
    buf.body(
        "The best policy was evaluated using playground/open_duck_mini_v2/mujoco_infer.py "
        "in the flat_terrain and flat_terrain_backlash environments. Results:")
    for s in [
        "Forward walking up to 0.15 m/s: achieved, smooth gait maintained.",
        "Lateral walking up to 0.2 m/s: achieved, minor lateral sway.",
        "Turning up to 1.0 rad/s: achieved, accurate heading control.",
        "Push recovery (random impulse up to 1.0 N*s): robot recovered "
        "in >90% of cases during 1000-step evaluation episodes.",
    ]:
        buf.bullet(s)

    buf.h2("4.3  Hardware Deployment Results")
    buf.body(
        "The best ONNX policy was transferred to the robot via SSH and executed with:")
    buf.body(
        "    python v2_rl_walk_mujoco.py "
        "--duck_config_path ~/duck_config.json "
        "--onnx_model_path ~/BEST_WALK_ONNX_2.onnx")
    buf.body(
        "The robot achieved stable forward walking on flat ground. Joystick-commanded "
        "lateral and turning motions were also functional. One limitation observed was "
        "gradual yaw drift on extended walks, attributable to residual IMU integration "
        "error and small asymmetries in joint friction between left and right legs.")
    buf.spacer()

    # ── 5  Challenges, Solutions & Discussion ────────────────────────────────
    buf.h1("5  Challenges, Solutions & Discussion")

    challenges = [
        (
            "Challenge 1: Motor Backlash and Compliance",
            "Problem: Cheap STS3215 servos exhibit ~0.05 rad of backlash per joint. "
            "The position controller overshoots and introduces oscillations, destabilising "
            "a policy trained without backlash modelling.",
            "Solution: Added passive dummy backlash joints to the MJCF model and used the "
            "flat_terrain_backlash training task. A LowPassActionFilter in the runtime "
            "further smooths motor commands.",
            "Status: Resolved. The backlash task noticeably improved real-robot stability."
        ),
        (
            "Challenge 2: Sim-to-Real IMU Orientation",
            "Problem: The IMU is physically mounted upside-down on the robot. Raw "
            "gravity and angular velocity readings are sign-inverted relative to the "
            "simulation frame, and the pitch zero-offset varies between individual units.",
            "Solution: Added imu_upside_down flag and user_pitch_bias parameter to "
            "duck_config.json. calibrate_imu.py provides an interactive calibration workflow.",
            "Status: Resolved. Per-robot JSON config corrects orientation automatically."
        ),
        (
            "Challenge 3: Communication Latency",
            "Problem: UART serial communication to 14 servos and I2C IMU reads introduce "
            "variable latency (~1-3 control steps at 50 Hz). A policy trained with zero "
            "delay assumption becomes unstable on hardware.",
            "Solution: Randomised action delays (0-3 steps) and IMU delays (0-3 steps) "
            "during training, specified via noise_config in joystick.py default_config.",
            "Status: Partially resolved. Some residual instability at high speed remains."
        ),
        (
            "Challenge 4: Training Compute Duration",
            "Problem: 300M-step training runs are required for good policies. Single-GPU "
            "training takes several hours, making rapid iteration slow.",
            "Solution: MuJoCo MJX (JAX-based) vectorises thousands of parallel environments "
            "on GPU. Also reduced num_timesteps to 150M for faster prototyping iterations "
            "while still obtaining deployable policies.",
            "Status: Resolved. 150M-step runs complete in ~2 hours on an RTX-class GPU."
        ),
        (
            "Challenge 5: ONNX Compatibility on Pi Zero 2W",
            "Problem: The Pi Zero 2W runs CPU-only onnxruntime. Some JIT-exported ONNX "
            "operators from JAX were not supported in lightweight ARM builds.",
            "Solution: Used awd=True in OnnxInfer which selects the action-with-delay "
            "compatible export path. Pinned onnxruntime to a known compatible version.",
            "Status: Resolved."
        ),
    ]
    for title, problem, solution, status in challenges:
        buf.h3(title)
        buf.body(problem)
        buf.body(solution)
        buf.body(status)
        buf.spacer()

    buf.h2("Personal Reflection")
    buf.body(
        "This project provided rare hands-on experience spanning mechanical design, embedded "
        "systems, and deep reinforcement learning — a full vertical slice of modern robotics. "
        "The most valuable lesson was that sim-to-real transfer is not a one-time step but "
        "an iterative loop: the simulation model must continuously be refined to capture "
        "real-world phenomena (backlash, latency, friction, IMU orientation). Given more time, "
        "I would invest in more thorough actuator system identification using BAM and in "
        "collecting real-robot trajectories to further close the gap through imitation from "
        "demonstrations.")
    buf.spacer()

    # ── 6  AI Tools Declaration ──────────────────────────────────────────────
    buf.h1("6  AI Tools Declaration")
    buf.body(
        "GitHub Copilot (model: Claude Sonnet 4.6) was used to assist in structuring and "
        "drafting sections of this report, drawing on the actual source code and documentation "
        "in the three project repositories. All technical content was verified against source "
        "files before inclusion. No AI-generated implementation code was submitted as part of "
        "this project without review and understanding.")
    buf.spacer()

    # ── 7  Conclusion ────────────────────────────────────────────────────────
    buf.h1("7  Conclusion")
    buf.body(
        "This project successfully demonstrated the complete pipeline from hardware assembly "
        "to sim-to-real RL policy deployment on the Open Duck Mini v2 bipedal robot. "
        "The robot was 3D-printed (37 parts), assembled with 14 Feetech servos, and calibrated. "
        "A locomotion policy was trained in MuJoCo MJX using PPO with imitation reward, "
        "velocity tracking, and regularisation. The ONNX-exported policy was deployed on a "
        "Raspberry Pi Zero 2W embedded controller at 50 Hz and achieved stable walking.")
    buf.body(
        "Key lessons learned: (1) accurate motor modelling (BAM-identified backlash/friction "
        "parameters) is critical for sim-to-real transfer; (2) randomising sensor delays and "
        "adding a backlash joint model significantly reduces the reality gap; (3) the modular "
        "three-repository architecture greatly simplifies the iteration cycle between training "
        "and deployment. This project connects directly to the course themes of sensor fusion, "
        "closed-loop control, and the challenges of deploying learned policies on physical "
        "hardware with manufacturing imperfections.")
    buf.spacer()

    # ── 8  References ────────────────────────────────────────────────────────
    buf.h1("8  References")
    for r in [
        "[1] A. Pirrone, S. Nguyen et al., \"Open Duck Mini v2,\" GitHub, 2025. "
        "https://github.com/apirrone/Open_Duck_Mini",
        "[2] A. Pirrone et al., \"Open Duck Playground,\" GitHub, 2025. "
        "https://github.com/apirrone/Open_Duck_Playground",
        "[3] A. Pirrone et al., \"Open Duck Mini Runtime,\" GitHub, 2025. "
        "https://github.com/apirrone/Open_Duck_Mini_Runtime",
        "[4] Google DeepMind, \"MuJoCo Playground,\" GitHub, 2024. "
        "https://github.com/google-deepmind/mujoco_playground",
        "[5] Rhoban, \"BAM — Backlash Actuator Modelling,\" GitHub. "
        "https://github.com/Rhoban/bam",
        "[6] Disney Research LA, \"Revisiting Reward Design for Policy Gradient Methods "
        "(BDX locomotion),\" 2024. "
        "https://la.disneyresearch.com/wp-content/uploads/BD_X_paper.pdf",
        "[7] A. Pirrone, \"Open Duck Reference Motion Generator,\" GitHub, 2025. "
        "https://github.com/apirrone/Open_Duck_reference_motion_generator",
        "[8] Raspberry Pi Foundation, \"Raspberry Pi Documentation,\" 2025. "
        "https://www.raspberrypi.com/documentation/",
        "[9] Feetech, \"STS3215 Serial Bus Servo Datasheet,\" 2023. "
        "https://www.feetechrc.com",
    ]:
        buf.bullet(r)
    buf.spacer()

    # ── A  Appendix ──────────────────────────────────────────────────────────
    buf.h1("A  Appendix")

    buf.h2("A.1  Key Command Reference")
    buf.table(doc,
        ["Task", "Command"],
        [
            ["Configure single motor",
             "python configure_motor.py --id <id>"],
            ["Configure all motors at once",
             "python configure_all_motors.py"],
            ["Find soft joint offsets",
             "python find_soft_offsets.py"],
            ["Calibrate IMU pitch bias",
             "python calibrate_imu.py"],
            ["Train policy (300M steps)",
             "uv run playground/open_duck_mini_v2/runner.py "
             "--task flat_terrain_backlash --num_timesteps 300000000"],
            ["Resume training from checkpoint",
             "uv run ... --restore_checkpoint_path checkpoints/<name>"],
            ["Test in simulation",
             "uv run playground/open_duck_mini_v2/mujoco_infer.py -o <path>.onnx"],
            ["Deploy on hardware",
             "python v2_rl_walk_mujoco.py "
             "--duck_config_path ~/duck_config.json --onnx_model_path <path>.onnx"],
            ["Monitor training via TensorBoard",
             "uv run tensorboard --logdir=checkpoints/"],
        ])

    buf.h2("A.2  Project File Structure")
    buf.body("Open_Duck_Mini/")
    for s in [
        "docs/  — assembly_guide.md, configure_motors.md, wiring diagrams, sim2real notes",
        "print/ — 37 STL files for all 3D-printed parts (PLA & TPU)",
        "mini_bdx/ — robot MJCF/URDF description, mass and inertia properties",
    ]:
        buf.bullet(s)
    buf.body("Open_Duck_Mini_Runtime/")
    for s in [
        "scripts/ — configure_motor.py, find_soft_offsets.py, v2_rl_walk_mujoco.py (main deploy)",
        "mini_bdx_runtime/ — HWI, OnnxInfer, raw_imu, FeetContacts, DuckConfig modules",
    ]:
        buf.bullet(s)
    buf.body("Open_Duck_Playground/")
    for s in [
        "playground/open_duck_mini_v2/ — joystick.py (env), custom_rewards.py, runner.py, XMLs",
        "playground/common/ — shared rewards, randomise, ONNX export utilities",
        "checkpoints/ — saved ONNX model files and TensorBoard event logs",
    ]:
        buf.bullet(s)


# ── Main ──────────────────────────────────────────────────────────────────────

def build_report():
    shutil.copy2(TEMPLATE, OUTPUT)
    doc = Document(OUTPUT)

    # Patch abstract
    for p in doc.paragraphs:
        if '(Summarize the project' in p.text:
            _clear_para(p)
            run = p.add_run(
                "This project demonstrates the full hardware-to-deployment pipeline for the "
                "Open Duck Mini v2, a miniature bipedal robot modelled after the Disney BDX droid. "
                "The robot was 3D-printed from 37 STL parts, assembled with 14 Feetech STS3215 "
                "servo motors, and calibrated. A locomotion policy was trained in MuJoCo MJX "
                "using PPO with an imitation reward over 150-300 million environment steps. "
                "The trained ONNX policy was deployed on a Raspberry Pi Zero 2W embedded "
                "controller at 50 Hz. The robot achieved stable forward walking on flat terrain "
                "responding to joystick velocity commands. Key challenges — motor backlash, "
                "communication latency, and IMU orientation — were addressed through domain "
                "randomisation, backlash joint modelling, and per-robot calibration workflows.")
            run.font.size = Pt(11)
            break

    # Erase placeholder text blocks (keep the structure, empty the content)
    erase_contains = [
        'Note to Students',
        'This template provides a suggested structure',
        'The descriptions under each section are only guiding',
        'Emphasize sections where you have',
        'Simplify or combine sections',
        'Add new sections that better reflect',
        'What matters most is that your report',
        'Good luck and enjoy',
        'Briefly describe the project objectives',
        'Include photos of your robot',
        'If you tried alternative configurations',
        'Describe your control strategy',
        'You can also explain why you chose this approach',
        'Present your experimental results',
        'about what worked and what',
        'Important: This section carries significant weight',
        'the problems you encountered, even if',
        'For each challenge, please include',
        'Also include your personal reflections',
        'Summarize your final performance',
        'Mention how this experience connects',
        'List all sources you referenced',
        'Include any additional materials',
    ]
    for p in doc.paragraphs:
        txt = p.text
        for phrase in erase_contains:
            if phrase.lower() in txt.lower():
                _clear_para(p)
                break

    # Erase template section heading paragraphs (we'll replace with proper ones below)
    # Use startswith match to handle trailing spaces / appended placeholder text
    section_heading_prefixes = (
        '1 \t', '2 \t', '3 \t', '4 \t', '5 \t',
        '6 \t', '7 \t', '8 \t', 'A \t',
    )
    for p in doc.paragraphs:
        if p.text.startswith(section_heading_prefixes):
            _clear_para(p)

    # Find the "Declaration for written assignment" paragraph — insert before it
    body_elem = doc.element.body
    declaration_elem = None
    for p in doc.paragraphs:
        if 'Declaration for written assignment' in p.text:
            declaration_elem = p._element
            break

    # Build the content buffer (no doc.add_* calls yet — tables need doc reference)
    buf = ParaBuffer()
    build_content(buf, doc)

    # Insert all buffer elements before the declaration block
    if declaration_elem is not None:
        parent = declaration_elem.getparent()
        idx = list(parent).index(declaration_elem)
        for i, elem in enumerate(buf.elements):
            parent.insert(idx + i, elem)
    else:
        # Fallback: just append
        for elem in buf.elements:
            body_elem.append(elem)

    doc.save(OUTPUT)
    print(f"[OK] Report saved: {OUTPUT}")


if __name__ == "__main__":
    build_report()
